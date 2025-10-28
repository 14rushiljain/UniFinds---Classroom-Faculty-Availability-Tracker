import re
from docx import Document


def parse_timetables_with_headers(file_stream):
    """
    Parses a .docx file stream, identifying tables based on preceding
    paragraph headers (e.g., 'Section-A').
    """
    doc = Document(file_stream)
    timetables = {}

    # Try to find headers like "Section-A"
    section_headers = [p.text.strip() for p in doc.paragraphs if
                       re.match(r'Section-[A-Z]', p.text.strip(), re.IGNORECASE)]

    # If headers don't match table count, use generic names
    if len(section_headers) != len(doc.tables):
        section_headers = [f"Section {i + 1}" for i in range(len(doc.tables))]

    for i, table in enumerate(doc.tables):
        section_name = section_headers[i]

        data = []
        current_day = ""
        for r, row in enumerate(table.rows):
            if r == 0: continue  # Skip header
            cells = [cell.text.strip() for cell in row.cells]
            if not any(cells): continue  # Skip empty rows

            if cells[0]:
                current_day = cells[0]
            else:
                cells[0] = current_day

            if len(cells) < 4: continue  # Skip malformed rows

            class_name = cells[2]
            subject_code = class_name.split(' ')[0]

            data.append({
                "Day": cells[0],
                "Time": cells[1],
                "Subject": f"{class_name} ({subject_code})",
                "Room": cells[3]
            })

        if data:
            timetables[section_name] = data

    return timetables
